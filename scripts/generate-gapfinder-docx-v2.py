"""
generate-gapfinder-docx.py

Fills a DOCX template using Phase 1 outputs, then optionally exports a PDF via Microsoft Word (docx2pdf).
"""

import os
import re
import csv
import json
from datetime import datetime
from collections import Counter, defaultdict

from openpyxl import load_workbook
from docx import Document
import importlib.util


ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
DATA_DIR = os.path.join(ROOT, "data")
TEMPLATE_PATH = os.path.join(ROOT, "templates", "gapfinder_readiness_template.docx")


def load_dotenv(root_dir: str):
    env_path = os.path.join(root_dir, ".env")
    if not os.path.exists(env_path):
        return

    with open(env_path, "r", encoding="utf-8") as f:
        for line in f:
            raw = line.strip()
            if not raw or raw.startswith("#") or "=" not in raw:
                continue

            key, value = raw.split("=", 1)
            key = key.strip()
            if not key or key in os.environ:
                continue

            value = value.strip()
            if (value.startswith('"') and value.endswith('"')) or (value.startswith("'") and value.endswith("'")):
                value = value[1:-1]

            os.environ[key] = value


load_dotenv(ROOT)


def clean_scope_path(pathname: str) -> str:
    p = (pathname or "/").strip()
    if not p.startswith("/"):
        p = f"/{p}"
    p = re.sub(r"/+$", "", p)
    return "" if p in ("", "/") else p


def audit_key_from_input(inp: str, args=None) -> str:
    args = args or []
    inp = (inp or "").strip()
    if not inp:
        return ""

    raw = inp if re.match(r"^https?://", inp, flags=re.I) else f"https://{inp}"

    from urllib.parse import urlparse
    parsed = urlparse(raw)
    host = re.sub(r"^www\.", "", parsed.hostname or "", flags=re.I)

    explicit_scope = None
    global_mode = "--global" in args or "--scope-mode=global" in args or "--scope-mode=none" in args
    for i, arg in enumerate(args):
        if arg.startswith("--scope-path="):
            explicit_scope = arg.split("=", 1)[1]
        elif arg == "--scope-path" and i + 1 < len(args):
            explicit_scope = args[i + 1]

    scope = "" if global_mode else clean_scope_path(parsed.path)
    if explicit_scope is not None:
        scope = clean_scope_path(explicit_scope)

    if not scope:
        return host

    suffix = "__".join(
        re.sub(r"[^a-zA-Z0-9._-]+", "-", part).strip("-")
        for part in scope.strip("/").split("/")
        if part
    )
    return f"{host}__{suffix}" if suffix else host


def safe_report_name(domain: str) -> str:
    return re.sub(r"[^a-zA-Z0-9._-]+", "_", domain)


def safe_sheet(wb, names):
    for n in names:
        if n in wb.sheetnames:
            return wb[n]
    return None


def header_map(sheet):
    if not sheet:
        return {}
    m = {}
    for cell in sheet[1]:
        if cell.value is not None:
            m[str(cell.value).strip()] = cell.col_idx
    return m


def cell_str(row, idx):
    if not idx:
        return ""
    v = row[idx - 1].value
    return str(v).strip() if v is not None else ""


def read_unknown_hosts(csv_path, top_n=10):
    if not os.path.exists(csv_path):
        return 0, []

    hosts = []
    with open(csv_path, "r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            h = (row.get("Host") or "").strip()
            if h:
                hosts.append(h)

    c = Counter(hosts)
    return len(hosts), [h for h, _ in c.most_common(top_n)]


def normalise_event_name(name: str) -> str:
    n = (name or "").strip()
    if not n:
        return ""

    n_low = n.lower().strip()
    mapping = {
        "pageview": "page_view",
        "page view": "page_view",
        "page_view": "page_view",
        "viewcontent": "view_item",
        "view_content": "view_item",
        "view item": "view_item",
        "view_item": "view_item",
        "addtocart": "add_to_cart",
        "add_to_cart": "add_to_cart",
        "begincheckout": "begin_checkout",
        "begin_checkout": "begin_checkout",
        "initiatecheckout": "begin_checkout",
        "purchase": "purchase",
    }
    return mapping.get(n_low, n)


def bullet_lines(items):
    return "\n".join([f"• {i}" for i in items if str(i).strip()])


def yesno(flag: bool) -> str:
    return "Yes" if flag else "Not observed in captured traffic"


def read_json_if_exists(p: str):
    if not os.path.exists(p):
        return None
    try:
        with open(p, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None


def fmt_val(v):
    if v is None or v == "":
        return "N/A"

    if isinstance(v, float):
        return f"{v:.3f}".rstrip("0").rstrip(".")

    return str(v)


def extract_vendor_presence(tag_sheet):
    if not tag_sheet:
        return {
            "has_gtm": False,
            "has_ga4": False,
            "ad_platforms": [],
            "vendor_count": 0,
            "top_by_category": {},
        }

    hm = header_map(tag_sheet)

    vendors = set()
    categories_seen = []
    ad_platforms = set()
    by_category = defaultdict(list)

    for r in tag_sheet.iter_rows(min_row=2, values_only=False):
        vendor = cell_str(r, hm.get("Vendor"))
        cat = cell_str(r, hm.get("Category"))

        if not vendor:
            continue

        vendors.add(vendor)

        if cat:
            categories_seen.append(cat)

        cat_l = (cat or "").lower()

        if "ads" in cat_l:
            if "Google Ads" in vendor:
                ad_platforms.add("Google Ads")
            if "Meta" in vendor:
                ad_platforms.add("Meta")
            if "TikTok" in vendor:
                ad_platforms.add("TikTok")
            if "Microsoft Ads" in vendor or "UET" in vendor:
                ad_platforms.add("Microsoft Ads")
            if "Pinterest" in vendor:
                ad_platforms.add("Pinterest")

        if cat and vendor not in by_category[cat]:
            by_category[cat].append(vendor)

    has_gtm = any("Tag Manager" in (c or "") for c in categories_seen) or any("Google Tag Manager" in v for v in vendors)
    has_ga4 = any(("Google Analytics" in v or "GA4" in v) for v in vendors)

    keep_cats = [
        "Analytics",
        "Ads",
        "Consent/CMP",
        "Session Replay",
        "A/B Testing",
        "Email/SMS",
        "Reviews/UGC",
        "Social Feed",
        "Support/Chat / Lead Capture",
        "Payments",
        "Search / Merchandising",
        "Server-side Tagging / Proxy",
    ]

    top_by_category = {}
    for k in keep_cats:
        vals = by_category.get(k, [])
        if not vals:
            for kk, vv in by_category.items():
                if k.lower() in (kk or "").lower():
                    vals = vv
                    break
        if vals:
            top_by_category[k] = ", ".join(vals[:3])

    return {
        "has_gtm": has_gtm,
        "has_ga4": has_ga4,
        "ad_platforms": sorted(list(ad_platforms)),
        "vendor_count": len(vendors),
        "top_by_category": top_by_category,
    }


def extract_event_stats(event_sheet):
    if not event_sheet:
        return {
            "event_count": 0,
            "top_events": [],
            "pct_value": None,
            "pct_currency": None,
            "pct_items": None,
            "pct_txn": None,
            "observed_event_set": set(),
        }

    hm = header_map(event_sheet)

    raw_events = []
    total = 0
    flags = {
        "HasValue": 0,
        "HasCurrency": 0,
        "HasItems": 0,
        "HasTransactionId": 0,
    }

    for r in event_sheet.iter_rows(min_row=2, values_only=False):
        total += 1

        ev = cell_str(r, hm.get("EventName"))
        if ev:
            raw_events.append(normalise_event_name(ev))

        for k in list(flags.keys()):
            val = cell_str(r, hm.get(k))
            if val.upper() == "Y":
                flags[k] += 1

    c = Counter([e for e in raw_events if e])
    top = [name for name, _ in c.most_common(6)]
    observed_set = set(c.keys())

    def pct(x):
        if total == 0:
            return None
        return round((x / total) * 100)

    return {
        "event_count": total,
        "top_events": top,
        "pct_value": pct(flags["HasValue"]),
        "pct_currency": pct(flags["HasCurrency"]),
        "pct_items": pct(flags["HasItems"]),
        "pct_txn": pct(flags["HasTransactionId"]),
        "observed_event_set": observed_set,
    }


def extract_privacy_summary(wb):
    sheet = safe_sheet(wb, [
        "privacy_summary",
        "baseline_privacy_summary",
        "probe_privacy_summary",
        "privacy_consent_summary",
    ])

    default = {
        "status": "Not evaluated",
        "score": "N/A",
        "cmp_observed": "Not observed",
        "consent_mode_observed": "Not observed",
        "consent_related_observed": "Not observed",
        "tracking_requests": "N/A",
        "evidence": [],
    }

    if not sheet:
        return default

    hm = header_map(sheet)

    for r in sheet.iter_rows(min_row=2, values_only=False):
        signal = cell_str(r, hm.get("Signal"))
        evidence = cell_str(r, hm.get("ExampleEvidence"))
        count = cell_str(r, hm.get("EvidenceCount"))

        if signal == "Privacy & Consent Visibility Status":
            default["status"] = evidence or "Not observed"
        elif signal == "Privacy & Consent Visibility Score":
            default["score"] = evidence or "N/A"
        elif signal == "CMP vendor observed":
            default["cmp_observed"] = evidence or "Not observed"
        elif signal == "Consent Mode signal observed":
            default["consent_mode_observed"] = evidence or "Not observed"
        elif signal == "Consent-related request observed":
            default["consent_related_observed"] = evidence or "Not observed"
        elif signal == "Tracking requests observed during scan":
            default["tracking_requests"] = count or "N/A"
        elif signal and evidence:
            default["evidence"].append(f"{signal}: {evidence}")

    return default


def build_privacy_consent_summary(privacy_info):
    lines = [
        f"Status: {privacy_info.get('status', 'Not evaluated')}",
        f"Visibility score: {privacy_info.get('score', 'N/A')}",
        "",
        "What this means:",
    ]

    if privacy_info.get("status") == "Consent signals observed":
        lines.append("Consent handling appears to be implemented at a basic level, but should be validated.")
    elif privacy_info.get("status") == "Partial visibility":
        lines.append("Consent signals are inconsistent or incomplete, which may affect data reliability.")
    else:
        lines.append("Consent handling was not clearly observed, which may introduce risk in data collection.")

    lines.append("")
    lines.append("This is a visibility scan only and does not confirm legal compliance.")

    return "\n".join(lines)

def build_privacy_consent_evidence(privacy_info):
    lines = []

    status = privacy_info.get("status")

    if status == "Consent signals observed":
        lines.append("Consent-related signals were detected across analytics and advertising requests.")
    elif status == "Partial visibility":
        lines.append("Some consent-related signals were detected, but coverage appears incomplete.")
    else:
        lines.append("No clear consent signals were observed during this scan.")

    if privacy_info.get("cmp_observed") == "Yes":
        lines.append("A recognised consent management platform is present.")
    else:
        lines.append("No recognised consent management platform was detected.")

    if privacy_info.get("consent_mode_observed") == "Yes":
        lines.append("Google Consent Mode-style behaviour was observed.")
    else:
        lines.append("Consent Mode signals were not clearly detected.")

    lines.append("This scan does not confirm whether consent choices are correctly enforced or compliant with regulations.")

    return bullet_lines(lines)

def build_tools_by_function(vendor_info):
    lines = []
    order = [
        "Analytics",
        "Ads",
        "Consent/CMP",
        "Email/SMS",
        "A/B Testing",
        "Session Replay",
        "Reviews/UGC",
        "Support/Chat / Lead Capture",
        "Search / Merchandising",
        "Payments",
        "Server-side Tagging / Proxy",
    ]

    for k in order:
        v = vendor_info["top_by_category"].get(k)
        if v:
            lines.append(f"{k}: {v}")

    return bullet_lines(lines) if lines else "• None observed"


def build_journey_signals(event_info):
    observed = event_info.get("observed_event_set", set())

    def has_any(candidates):
        return any(c in observed for c in candidates)

    browsing = has_any({"page_view", "view_item", "view_item_list"})
    product_interest = has_any({"view_item", "select_item", "view_content"})
    cart = has_any({"add_to_cart"})
    checkout = has_any({"begin_checkout", "add_shipping_info", "add_payment_info"})
    purchase = has_any({"purchase"})

    lines = [
        f"Browsing signals (e.g. page/product views): {yesno(browsing)}",
        f"Product interest signals (e.g. view_item): {yesno(product_interest)}",
        f"Cart signals (e.g. add_to_cart): {yesno(cart)}",
        f"Checkout signals (e.g. begin_checkout): {yesno(checkout)}",
        f"Purchase signals (e.g. purchase): {yesno(purchase)}",
    ]

    return bullet_lines(lines)


def build_payload_completeness(event_info):
    pv = event_info.get("pct_value")
    pc = event_info.get("pct_currency")
    pi = event_info.get("pct_items")
    pt = event_info.get("pct_txn")

    def fmt(p):
        return "N/A" if p is None else f"{p}%"

    lines = [
        f"Value present (observed): {fmt(pv)}",
        f"Currency present (observed): {fmt(pc)}",
        f"Items present (observed): {fmt(pi)}",
        f"Transaction ID present (observed): {fmt(pt)}",
    ]

    return bullet_lines(lines)


def build_event_reference(event_info):
    observed = event_info.get("observed_event_set", set())

    expected_events = [
        ("page_view", "Browsing visibility"),
        ("view_item", "Product interest visibility"),
        ("add_to_cart", "Cart behaviour visibility"),
        ("begin_checkout", "Checkout intent visibility"),
        ("purchase", "Revenue / purchase visibility"),
    ]

    lines = []

    for event_name, meaning in expected_events:
        status = "Observed" if event_name in observed else "Not observed"
        lines.append(f"{event_name}: {status} — {meaning}")

    return bullet_lines(lines)


def build_payload_reference(event_info):
    pv = event_info.get("pct_value")
    pc = event_info.get("pct_currency")
    pi = event_info.get("pct_items")
    pt = event_info.get("pct_txn")

    def fmt(p):
        return "N/A" if p is None else f"{p}%"

    lines = [
        f"value: {fmt(pv)} observed",
        f"currency: {fmt(pc)} observed",
        f"items/product data: {fmt(pi)} observed",
        f"transaction_id: {fmt(pt)} observed",
    ]

    return bullet_lines(lines)


def build_coverage_summary(domain, vendor_info, event_info):
    platforms = ", ".join(vendor_info["ad_platforms"]) if vendor_info["ad_platforms"] else "Not observed in captured traffic"

    lines = [
        f"Domain: {domain}",
        f"Tracking foundation: GTM {yesno(vendor_info['has_gtm'])} • GA4 {yesno(vendor_info['has_ga4'])}",
        f"Paid platforms observed: {platforms}",
        f"Distinct vendors observed: {vendor_info['vendor_count']}",
        f"Total events observed: {event_info.get('event_count', 0)}",
    ]

    return bullet_lines(lines)


def build_attribution_summary():
    lines = [
        "Campaign parameters (UTMs): Not evaluated in this run (probe required)",
        "Platform click IDs (gclid/fbclid/ttclid/wbraid): Not observed in captured traffic",
        "Attribution sensitivity: Reported performance can shift based on configuration choices",
    ]

    return bullet_lines(lines)


def build_unknown_summary(unknown_count, unknown_top):
    if unknown_count == 0:
        return bullet_lines([
            "Unknown vendors observed: 0",
            "Top unknown hosts: None",
        ])

    return bullet_lines([
        f"Unknown vendors observed: {unknown_count}",
        f"Top unknown hosts: {', '.join(unknown_top) if unknown_top else 'None'}",
    ])


def apply_psi(mapping: dict, psi: dict):
    home = (psi or {}).get("targets", {}).get("home", {}) if psi else {}
    m = (home or {}).get("mobile", {})
    d = (home or {}).get("desktop", {})

    mapping.update({
        "MobileCLS": fmt_val(m.get("cls")),
        "MobileLCP": fmt_val(m.get("lcp_s")),
        "MobileFCP": fmt_val(m.get("fcp_s")),
        "MobileTTI": fmt_val(m.get("tti_s")),
        "MobileTBT": fmt_val(m.get("tbt_ms")),
        "MobileOverall": fmt_val(m.get("performance")),

        "DesktopCLS": fmt_val(d.get("cls")),
        "DesktopLCP": fmt_val(d.get("lcp_s")),
        "DesktopFCP": fmt_val(d.get("fcp_s")),
        "DesktopTTI": fmt_val(d.get("tti_s")),
        "DesktopTBT": fmt_val(d.get("tbt_ms")),
        "DesktopOverall": fmt_val(d.get("performance")),
    })


def score_tracking_foundation(vendor_info):
    score = 0
    reasons = []

    if vendor_info.get("has_gtm"):
        score += 1
    else:
        reasons.append("GTM was not observed.")

    if vendor_info.get("has_ga4"):
        score += 1
    else:
        reasons.append("GA4 was not observed.")

    return score, reasons


def score_event_visibility(event_info):
    observed = event_info.get("observed_event_set", set())

    product_events = {"view_item", "select_item", "view_content"}
    cart_events = {"add_to_cart"}
    checkout_events = {"begin_checkout", "add_shipping_info", "add_payment_info"}
    purchase_events = {"purchase"}

    has_product = bool(observed.intersection(product_events))
    has_cart = bool(observed.intersection(cart_events))
    has_checkout = bool(observed.intersection(checkout_events))
    has_purchase = bool(observed.intersection(purchase_events))

    funnel_points = sum([has_product, has_cart, has_checkout, has_purchase])

    if has_purchase or has_checkout:
        score = 2
    elif funnel_points >= 1:
        score = 1
    else:
        score = 0

    reasons = []

    if not has_product:
        reasons.append("Product-level engagement signals were not observed.")
    if not has_cart:
        reasons.append("Cart signals were not observed.")
    if not has_checkout:
        reasons.append("Checkout signals were not observed.")
    if not has_purchase:
        reasons.append("Purchase signals were not observed.")

    return score, reasons


def score_payload_quality(event_info):
    value = event_info.get("pct_value") or 0
    currency = event_info.get("pct_currency") or 0
    items = event_info.get("pct_items") or 0
    txn = event_info.get("pct_txn") or 0

    strong = value >= 80 and currency >= 80 and txn >= 80
    partial = value > 0 or currency > 0 or items > 0 or txn > 0

    if strong:
        score = 2
    elif partial:
        score = 1
    else:
        score = 0

    reasons = []

    if value < 80:
        reasons.append(f"Value data is only present on {value}% of observed events.")
    if currency < 80:
        reasons.append(f"Currency data is only present on {currency}% of observed events.")
    if items < 80:
        reasons.append(f"Item/product data is only present on {items}% of observed events.")
    if txn < 80:
        reasons.append(f"Transaction IDs are only present on {txn}% of observed events.")

    return score, reasons


def score_performance(psi):
    home = (psi or {}).get("targets", {}).get("home", {}) if psi else {}
    mobile = (home or {}).get("mobile", {})
    desktop = (home or {}).get("desktop", {})

    mobile_score = mobile.get("performance")
    desktop_score = desktop.get("performance")

    if mobile_score is None:
        return 0, ["Mobile performance could not be measured."]

    if mobile_score >= 90:
        score = 2
    elif mobile_score >= 70:
        score = 1.5
    elif mobile_score >= 50:
        score = 1
    elif mobile_score >= 30:
        score = 0.5
    else:
        score = 0

    reasons = []

    if mobile_score < 70:
        reasons.append(f"Mobile performance is weak at {mobile_score}/100.")
    if desktop_score is not None and desktop_score < 90:
        reasons.append(f"Desktop performance is below ideal at {desktop_score}/100.")

    return score, reasons


def score_attribution_readiness():
    return 0, [
        "UTM persistence was not confirmed.",
        "Platform click IDs were not observed in captured traffic.",
    ]


def score_gapfinder(vendor_info, event_info, psi):
    tracking_score, tracking_reasons = score_tracking_foundation(vendor_info)
    event_score, event_reasons = score_event_visibility(event_info)
    payload_score, payload_reasons = score_payload_quality(event_info)
    performance_score, performance_reasons = score_performance(psi)
    attribution_score, attribution_reasons = score_attribution_readiness()

    total = tracking_score + event_score + payload_score + performance_score + attribution_score

    if total <= 3:
        status = "At Risk"
    elif total <= 6:
        status = "Underperforming"
    elif total <= 8:
        status = "Strong"
    else:
        status = "Optimised"

    all_reasons = (
        event_reasons +
        payload_reasons +
        performance_reasons +
        attribution_reasons +
        tracking_reasons
    )

    top_issues = all_reasons[:4] if all_reasons else ["No major structural issues detected from this scan."]

    return {
        "overall_score": round(total, 1),
        "status": status,
        "tracking_score": tracking_score,
        "event_score": event_score,
        "payload_score": payload_score,
        "performance_score": performance_score,
        "attribution_score": attribution_score,
        "top_issues": top_issues,
    }


def build_top_issues(score_info):
    return bullet_lines(score_info.get("top_issues", []))


def status_from_overall(score):
    if score is None:
        return "Not scored"
    if score < 40:
        return "At Risk"
    if score < 60:
        return "Underperforming"
    if score < 80:
        return "Strong"
    return "Optimised"


def score_info_from_scorecard(scorecard):
    categories = (scorecard or {}).get("categories", {})
    issues = (scorecard or {}).get("criticalIssues", [])

    def c(name):
        return round(float(categories.get(name, {}).get("capped_score", 0)), 1)

    top_issues = [i.get("recommendation") or i.get("code") for i in issues[:4] if isinstance(i, dict)]
    overall = (scorecard or {}).get("overallScore")

    return {
        "overall_score": round(float(overall), 1) if overall is not None else 0.0,
        "status": status_from_overall(overall),
        "tracking_score": c("tracking_foundation"),
        "event_score": c("event_signal_integrity"),
        "payload_score": c("ecommerce_signal_quality"),
        "performance_score": c("performance_friction"),
        "attribution_score": c("platform_signal_alignment"),
        "top_issues": top_issues or ["No major structural issues detected from this scan."],
    }


def _replace_in_paragraph(paragraph, mapping: dict):
    if not paragraph.runs:
        return

    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text

    for key, val in mapping.items():
        token = f"{{{{{key}}}}}"
        if token in new_text:
            new_text = new_text.replace(token, str(val))

    if new_text == full_text:
        return

    paragraph.runs[0].text = new_text
    for r in paragraph.runs[1:]:
        r.text = ""


def _replace_in_cell(cell, mapping: dict):
    for p in cell.paragraphs:
        _replace_in_paragraph(p, mapping)


def replace_placeholders_in_doc(doc: Document, mapping: dict):
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                _replace_in_cell(cell, mapping)

    for section in doc.sections:
        for p in section.header.paragraphs:
            _replace_in_paragraph(p, mapping)
        for t in section.header.tables:
            for row in t.rows:
                for cell in row.cells:
                    _replace_in_cell(cell, mapping)

        for p in section.footer.paragraphs:
            _replace_in_paragraph(p, mapping)
        for t in section.footer.tables:
            for row in t.rows:
                for cell in row.cells:
                    _replace_in_cell(cell, mapping)



def pdf_requested(args=None) -> bool:
    args = args or []
    if "--no-pdf" in args:
        return False
    if "--pdf" in args:
        return True

    v = os.environ.get("GAPFINDER_EXPORT_PDF", "").strip().lower()
    if v in {"0", "false", "no", "off", "skip"}:
        return False
    if v in {"1", "true", "yes", "on"}:
        return True

    return True


def pdf_required(args=None) -> bool:
    args = args or []
    if "--require-pdf" in args:
        return True

    v = os.environ.get("GAPFINDER_REQUIRE_PDF", "").strip().lower()
    return v in {"1", "true", "yes", "on"}


def docx2pdf_available() -> bool:
    return importlib.util.find_spec("docx2pdf") is not None


def export_pdf_if_requested(out_docx: str, out_pdf: str, args=None):
    if not pdf_requested(args):
        print("[Skip] PDF export disabled (--no-pdf or GAPFINDER_EXPORT_PDF=false).")
        return False

    if not docx2pdf_available():
        msg = "docx2pdf is not installed. DOCX report was created; skipping PDF export."
        if pdf_required(args):
            raise RuntimeError(msg)
        print(f"[WARN] {msg}")
        return False

    from docx2pdf import convert

    try:
        convert(out_docx, out_pdf)
        return True
    except Exception as e:
        msg = (
            "DOCX report was created, but PDF conversion failed. "
            "On macOS, docx2pdf requires Microsoft Word and may fail if Word is missing, "
            "not permitted by macOS Automation permissions, or the document is open. "
            "Use --no-pdf or GAPFINDER_EXPORT_PDF=false to skip PDF export.\n"
            f"Details: {e}"
        )
        if pdf_required(args):
            raise RuntimeError(msg)
        print(f"[WARN] {msg}")
        return False


def add_paragraphs(doc: Document, lines):
    for line in lines:
        if not str(line).strip():
            doc.add_paragraph("")
        elif str(line).startswith("# "):
            doc.add_heading(str(line)[2:], level=1)
        elif str(line).startswith("## "):
            doc.add_heading(str(line)[3:], level=2)
        else:
            doc.add_paragraph(str(line))


def build_fallback_doc(mapping: dict) -> Document:
    doc = Document()
    doc.add_heading("GapFinder Readiness Report", level=0)
    add_paragraphs(doc, [
        f"Website: {mapping.get('website', 'N/A')}",
        f"Generated: {mapping.get('GeneratedAt', 'N/A')}",
        "",
        "## Scorecard",
        f"Overall score: {mapping.get('GapFinderScore', 'N/A')}",
        f"Status: {mapping.get('GapFinderStatus', 'N/A')}",
        f"Tracking foundation: {mapping.get('TrackingFoundationScore', 'N/A')}",
        f"Event visibility: {mapping.get('EventVisibilityScore', 'N/A')}",
        f"Payload quality: {mapping.get('PayloadQualityScore', 'N/A')}",
        f"Performance: {mapping.get('PerformanceScore', 'N/A')}",
        f"Attribution readiness: {mapping.get('AttributionReadinessScore', 'N/A')}",
        "",
        "## Top Issues",
        mapping.get("TopIssues", "None observed"),
        "",
        "## Tracking Coverage",
        mapping.get("CoverageSummary", "N/A"),
        "",
        "## Vendors by Function",
        mapping.get("ToolsByFunction", "N/A"),
        "",
        "## Journey Signals",
        mapping.get("JourneySignals", "N/A"),
        "",
        "## Payload Completeness",
        mapping.get("PayloadCompleteness", "N/A"),
        "",
        "## Privacy & Consent",
        mapping.get("PrivacyConsentSummary", "N/A"),
        "",
        "## Unknown Vendors",
        mapping.get("UnknownSummary", "N/A"),
    ])
    return doc


def load_report_document(mapping: dict) -> Document:
    if os.path.exists(TEMPLATE_PATH):
        doc = Document(TEMPLATE_PATH)
        replace_placeholders_in_doc(doc, mapping)
        return doc

    print(f"[WARN] Template not found: {TEMPLATE_PATH}")
    print("[WARN] Creating a fallback DOCX report instead. Restore templates/gapfinder_readiness_template.docx for branded output.")
    return build_fallback_doc(mapping)

def main(domain_input: str, args=None):
    domain = audit_key_from_input(domain_input, args or [])

    analysis_dir = os.path.join(DATA_DIR, domain, "analysis")
    xlsx_path = os.path.join(analysis_dir, "phase1_inventory.xlsx")
    unknown_path = os.path.join(analysis_dir, "unknown_vendors.csv")
    psi_path = os.path.join(analysis_dir, "psi.json")
    scorecard_path = os.path.join(analysis_dir, "scorecard.json")

    if not os.path.exists(xlsx_path):
        raise FileNotFoundError(f"Missing workbook: {xlsx_path}")

    wb = load_workbook(xlsx_path, data_only=True)

    tag_sheet = safe_sheet(wb, ["baseline_tag_inventory", "tag_inventory", "Tag Inventory"])
    event_sheet = safe_sheet(wb, ["baseline_event_inventory", "event_inventory", "Event Inventory"])

    vendor_info = extract_vendor_presence(tag_sheet)
    event_info = extract_event_stats(event_sheet)
    privacy_info = extract_privacy_summary(wb)
    unknown_count, unknown_top = read_unknown_hosts(unknown_path, top_n=10)
    psi = read_json_if_exists(psi_path)
    scorecard = read_json_if_exists(scorecard_path)

    score_info = score_info_from_scorecard(scorecard) if scorecard else score_gapfinder(vendor_info, event_info, psi)

    mapping = {
        "website": domain,
        "GeneratedAt": datetime.now().strftime("%Y-%m-%d %H:%M"),

        "GapFinderScore": str(score_info["overall_score"]),
        "GapFinderStatus": score_info["status"],
        "TrackingFoundationScore": str(score_info["tracking_score"]),
        "EventVisibilityScore": str(score_info["event_score"]),
        "PayloadQualityScore": str(score_info["payload_score"]),
        "PerformanceScore": str(score_info["performance_score"]),
        "AttributionReadinessScore": str(score_info["attribution_score"]),
        "TopIssues": build_top_issues(score_info),

        "HasGTM": "Yes" if vendor_info["has_gtm"] else "Not observed in captured traffic",
        "HasGA4": "Yes" if vendor_info["has_ga4"] else "Not observed in captured traffic",
        "AdPlatforms": ", ".join(vendor_info["ad_platforms"]) if vendor_info["ad_platforms"] else "Not observed in captured traffic",
        "VendorCount": str(vendor_info["vendor_count"]),
        "TopEvents": ", ".join(event_info["top_events"]) if event_info["top_events"] else "None observed",
        "EventCount": str(event_info["event_count"]),
        "EventReference": build_event_reference(event_info),
        "PayloadReference": build_payload_reference(event_info),
        "PctHasValue": "N/A" if event_info["pct_value"] is None else f"{event_info['pct_value']}%",
        "PctHasCurrency": "N/A" if event_info["pct_currency"] is None else f"{event_info['pct_currency']}%",
        "PctHasItems": "N/A" if event_info["pct_items"] is None else f"{event_info['pct_items']}%",
        "PctHasTransactionId": "N/A" if event_info["pct_txn"] is None else f"{event_info['pct_txn']}%",
        "UTMsObserved": "Inconclusive from website signals alone",
        "UTMsPersist": "Inconclusive from website signals alone",
        "ClickIdsObserved": "Not observed in captured traffic",
        "UnknownCount": str(unknown_count),
        "TopUnknownHosts": ", ".join(unknown_top) if unknown_top else "None",

        "PrivacyConsentStatus": privacy_info.get("status", "Not evaluated"),
        "PrivacyConsentScore": privacy_info.get("score", "N/A"),
        "PrivacyConsentSummary": build_privacy_consent_summary(privacy_info),
        "PrivacyConsentEvidence": build_privacy_consent_evidence(privacy_info),

        "MobileCLS": "N/A",
        "MobileLCP": "N/A",
        "MobileFCP": "N/A",
        "MobileTTI": "N/A",
        "MobileTBT": "N/A",
        "MobileOverall": "N/A",
        "DesktopCLS": "N/A",
        "DesktopLCP": "N/A",
        "DesktopFCP": "N/A",
        "DesktopTTI": "N/A",
        "DesktopTBT": "N/A",
        "DesktopOverall": "N/A",
    }

    mapping["CoverageSummary"] = build_coverage_summary(domain, vendor_info, event_info)
    mapping["ToolsByFunction"] = build_tools_by_function(vendor_info)
    mapping["JourneySignals"] = build_journey_signals(event_info)
    mapping["PayloadCompleteness"] = build_payload_completeness(event_info)
    mapping["AttributionSummary"] = build_attribution_summary()
    mapping["UnknownSummary"] = build_unknown_summary(unknown_count, unknown_top)

    parts = []
    for cat, vendors in vendor_info["top_by_category"].items():
        parts.append(f"{cat}: {vendors}")

    mapping["TopVendorsByCategory"] = " | ".join(parts) if parts else "None observed"

    if psi:
        apply_psi(mapping, psi)

    doc = load_report_document(mapping)

    report_dir = os.path.join(DATA_DIR, domain, "report")
    os.makedirs(report_dir, exist_ok=True)

    out_docx = os.path.join(report_dir, f"GapFinder_Readiness_{safe_report_name(domain)}.docx")
    doc.save(out_docx)

    out_pdf = os.path.splitext(out_docx)[0] + ".pdf"

    pdf_created = export_pdf_if_requested(out_docx, out_pdf, args or [])

    print(f"[OK] Wrote DOCX: {out_docx}")
    if pdf_created:
        print(f"[OK] Wrote PDF:  {out_pdf}")
    else:
        print("[OK] PDF export skipped or unavailable; DOCX report is ready.")


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python scripts/generate-gapfinder-docx-v2.py <domain> [--no-pdf] [--require-pdf]")
        raise SystemExit(1)

    main(sys.argv[1], sys.argv[2:])