"""
generate-gapfinder-docx.py

Fills a DOCX template using Phase 1 outputs, then exports a PDF via Microsoft Word (docx2pdf).

Usage (run from project root):
  python scripts/generate-gapfinder-docx.py latexmattress.com.au

Assumes template exists at:
  templates/gapfinder_readiness_template.docx

Inputs:
  data/<domain>/analysis/phase1_inventory.xlsx
  data/<domain>/analysis/unknown_vendors.csv

Outputs:
  data/<domain>/report/GapFinder_Readiness_<domain>.docx
  data/<domain>/report/GapFinder_Readiness_<domain>.pdf
"""

import os
import re
import csv
from datetime import datetime
from collections import Counter, defaultdict

from openpyxl import load_workbook
from docx import Document
from docx2pdf import convert


# -----------------------------
# Paths / constants
# -----------------------------
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
DATA_DIR = os.path.join(ROOT, "data")
TEMPLATE_PATH = os.path.join(ROOT, "templates", "gapfinder_readiness_template.docx")


# -----------------------------
# Utilities
# -----------------------------
def normalise_domain(inp: str) -> str:
    inp = (inp or "").strip()
    inp = re.sub(r"^https?://", "", inp, flags=re.I)
    inp = inp.split("/")[0]
    return inp.replace("www.", "")

def safe_sheet(wb, names):
    for n in names:
        if n in wb.sheetnames:
            return wb[n]
    return None

def header_map(sheet):
    if not sheet:
        return {}
    m = {}
    for cell in sheet[1]:
        if cell.value is None:
            continue
        m[str(cell.value).strip()] = cell.col_idx
    return m

def cell_str(row, idx):
    if not idx:
        return ""
    v = row[idx - 1].value
    return str(v).strip() if v is not None else ""

def read_unknown_hosts(csv_path, top_n=10):
    if not os.path.exists(csv_path):
        return 0, []
    hosts = []
    with open(csv_path, "r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            h = (row.get("Host") or "").strip()
            if h:
                hosts.append(h)
    c = Counter(hosts)
    return len(hosts), [h for h, _ in c.most_common(top_n)]


# -----------------------------
# Extraction logic
# -----------------------------
def extract_vendor_presence(tag_sheet):
    """
    Expects columns like:
      Vendor, Category, Host, Identifier, FirstParty, RequestCount
    """
    if not tag_sheet:
        return {
            "has_gtm": "Not observed in captured traffic",
            "has_ga4": "Not observed in captured traffic",
            "ad_platforms": "Not observed in captured traffic",
            "vendor_count": "0",
            "top_by_category": {}
        }

    hm = header_map(tag_sheet)

    vendors = set()
    categories_seen = []
    ad_platforms = set()
    by_category = defaultdict(list)

    for r in tag_sheet.iter_rows(min_row=2, values_only=False):
        vendor = cell_str(r, hm.get("Vendor"))
        cat = cell_str(r, hm.get("Category"))

        if not vendor:
            continue

        vendors.add(vendor)
        if cat:
            categories_seen.append(cat)

        cat_l = (cat or "").lower()

        # recognisable ad platforms only (layman-safe)
        if "ads" in cat_l:
            if "Google Ads" in vendor: ad_platforms.add("Google Ads")
            if "Meta" in vendor: ad_platforms.add("Meta")
            if "TikTok" in vendor: ad_platforms.add("TikTok")
            if "Microsoft Ads" in vendor or "UET" in vendor: ad_platforms.add("Microsoft Ads")
            if "Pinterest" in vendor: ad_platforms.add("Pinterest")

        # cap per category
        if cat and vendor not in by_category[cat]:
            by_category[cat].append(vendor)

    has_gtm = "Yes" if any("Tag Manager" in (c or "") for c in categories_seen) or any("Google Tag Manager" in v for v in vendors) else "Not observed in captured traffic"
    has_ga4 = "Yes" if any(("Google Analytics" in v or "GA4" in v) for v in vendors) else "Not observed in captured traffic"
    ad_platforms_str = ", ".join(sorted(ad_platforms)) if ad_platforms else "Not observed in captured traffic"

    # Common categories in your ruleset — keep short, credible, and readable
    keep_cats = [
        "Analytics",
        "Ads",
        "Consent/CMP",
        "Session Replay",
        "A/B Testing",
        "Email/SMS",
        "Reviews/UGC",
        "Social Feed",
        "Support/Chat / Lead Capture",
        "Payments",
        "Search / Merchandising",
        "Server-side Tagging / Proxy",
    ]

    top_by_category = {}
    for k in keep_cats:
        vals = by_category.get(k, [])
        if not vals:
            # fuzzy match
            for kk, vv in by_category.items():
                if k.lower() in (kk or "").lower():
                    vals = vv
                    break
        if vals:
            top_by_category[k] = ", ".join(vals[:3])

    return {
        "has_gtm": has_gtm,
        "has_ga4": has_ga4,
        "ad_platforms": ad_platforms_str,
        "vendor_count": str(len(vendors)),
        "top_by_category": top_by_category
    }

def extract_event_stats(event_sheet):
    """
    Expects columns like:
      EventName, HasValue, HasCurrency, HasItems, HasTransactionId
    """
    if not event_sheet:
        return {
            "event_count": "0",
            "top_events": "None observed",
            "pct_value": "N/A",
            "pct_currency": "N/A",
            "pct_items": "N/A",
            "pct_txn": "N/A"
        }

    hm = header_map(event_sheet)

    events = []
    total = 0
    flags = {"HasValue": 0, "HasCurrency": 0, "HasItems": 0, "HasTransactionId": 0}

    for r in event_sheet.iter_rows(min_row=2, values_only=False):
        total += 1
        ev = cell_str(r, hm.get("EventName"))
        if ev:
            events.append(ev)

        for k in list(flags.keys()):
            val = cell_str(r, hm.get(k))
            if val.upper() == "Y":
                flags[k] += 1

    top = ", ".join([name for name, _ in Counter(events).most_common(6)]) if events else "None observed"

    def pct(x):
        if total == 0:
            return "N/A"
        return f"{round((x/total)*100)}%"

    return {
        "event_count": str(total),
        "top_events": top,
        "pct_value": pct(flags["HasValue"]),
        "pct_currency": pct(flags["HasCurrency"]),
        "pct_items": pct(flags["HasItems"]),
        "pct_txn": pct(flags["HasTransactionId"])
    }

def extract_probe_flags(wb):
    """
    Conservative defaults: we only claim what the workbook explicitly exposes.
    If you later add explicit fields to probe_summary/delta_summary tabs, wire them in here.
    """
    # If you have these sheets, great; but we won't scrape arbitrary cells blindly.
    _ = safe_sheet(wb, ["probe_summary", "Probe Summary"])
    _ = safe_sheet(wb, ["delta_summary", "Delta Summary"])

    return {
        "utms_observed": "Inconclusive from website signals alone",
        "utms_persist": "Inconclusive from website signals alone",
        "click_ids": "Not observed in captured traffic",
    }


# -----------------------------
# DOCX placeholder replacement
# -----------------------------
def _replace_in_paragraph(paragraph, mapping: dict):
    """
    Replace placeholders even if Word split them across runs.
    Strategy:
      1) Join all runs to a single string
      2) Replace tokens in the full string
      3) Write back into runs (preserve formatting as best-effort by putting
         all text into the first run and clearing the rest)
    """
    if not paragraph.runs:
        return

    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text

    for key, val in mapping.items():
        token = f"{{{{{key}}}}}"
        if token in new_text:
            new_text = new_text.replace(token, str(val))

    # No changes
    if new_text == full_text:
        return

    # Put all text into first run; clear the others
    paragraph.runs[0].text = new_text
    for r in paragraph.runs[1:]:
        r.text = ""


def _replace_in_cell(cell, mapping: dict):
    for p in cell.paragraphs:
        _replace_in_paragraph(p, mapping)


def replace_placeholders_in_doc(doc: Document, mapping: dict):
    # Body paragraphs
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)

    # Tables
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                _replace_in_cell(cell, mapping)

    # Headers / footers (common place for templates)
    for section in doc.sections:
        header = section.header
        footer = section.footer

        for p in header.paragraphs:
            _replace_in_paragraph(p, mapping)
        for t in header.tables:
            for row in t.rows:
                for cell in row.cells:
                    _replace_in_cell(cell, mapping)

        for p in footer.paragraphs:
            _replace_in_paragraph(p, mapping)
        for t in footer.tables:
            for row in t.rows:
                for cell in row.cells:
                    _replace_in_cell(cell, mapping)



# -----------------------------
# Main
# -----------------------------
def main(domain_input: str):
    domain = normalise_domain(domain_input)

    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(f"Template not found: {TEMPLATE_PATH}")

    analysis_dir = os.path.join(DATA_DIR, domain, "analysis")
    xlsx_path = os.path.join(analysis_dir, "phase1_inventory.xlsx")
    unknown_path = os.path.join(analysis_dir, "unknown_vendors.csv")

    if not os.path.exists(xlsx_path):
        raise FileNotFoundError(f"Missing workbook: {xlsx_path}")

    wb = load_workbook(xlsx_path, data_only=True)

    tag_sheet = safe_sheet(wb, ["baseline_tag_inventory", "tag_inventory", "Tag Inventory"])
    event_sheet = safe_sheet(wb, ["baseline_event_inventory", "event_inventory", "Event Inventory"])

    vendor_info = extract_vendor_presence(tag_sheet)
    event_info = extract_event_stats(event_sheet)
    unknown_count, unknown_top = read_unknown_hosts(unknown_path, top_n=10)
    probe_info = extract_probe_flags(wb)

    # Map placeholders to values used in your template
    mapping = {
        "website": domain,
        "GeneratedAt": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "HasGTM": vendor_info["has_gtm"],
        "HasGA4": vendor_info["has_ga4"],
        "AdPlatforms": vendor_info["ad_platforms"],
        "VendorCount": vendor_info["vendor_count"],
        "TopEvents": event_info["top_events"],
        "EventCount": event_info["event_count"],
        "PctHasValue": event_info["pct_value"],
        "PctHasCurrency": event_info["pct_currency"],
        "PctHasItems": event_info["pct_items"],
        "PctHasTransactionId": event_info["pct_txn"],
        "UTMsObserved": probe_info["utms_observed"],
        "UTMsPersist": probe_info["utms_persist"],
        "ClickIdsObserved": probe_info["click_ids"],
        "UnknownCount": str(unknown_count),
        "TopUnknownHosts": ", ".join(unknown_top) if unknown_top else "None",
    }

    # Optional: allow template to include category callouts like {{Tools_Analytics}}
    for k, v in vendor_info["top_by_category"].items():
        safe_key = "Tools_" + re.sub(r"[^A-Za-z0-9]+", "_", k).strip("_")
        mapping[safe_key] = v

    # Load template, fill placeholders
    doc = Document(TEMPLATE_PATH)
    replace_placeholders_in_doc(doc, mapping)

    # Save outputs
    report_dir = os.path.join(DATA_DIR, domain, "report")
    os.makedirs(report_dir, exist_ok=True)

    out_docx = os.path.join(report_dir, f"GapFinder_Readiness_{domain}.docx")
    doc.save(out_docx)

    # Convert to PDF via Word
    out_pdf = os.path.splitext(out_docx)[0] + ".pdf"
    try:
        convert(out_docx, out_pdf)
    except Exception as e:
        raise RuntimeError(
            "DOCX saved, but PDF conversion failed. "
            "Make sure Microsoft Word is installed and the DOCX is not open.\n"
            f"Details: {e}"
        )

    print(f"[OK] Wrote DOCX: {out_docx}")
    print(f"[OK] Wrote PDF:  {out_pdf}")


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python scripts/generate-gapfinder-docx.py <domain>")
        raise SystemExit(1)
    main(sys.argv[1])